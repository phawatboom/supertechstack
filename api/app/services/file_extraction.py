from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader


class FileExtractionError(Exception):
    """Base exception for file extraction failures."""


class UnsupportedFileTypeError(FileExtractionError):
    """Raised when an uploaded file extension is unsupported."""


class InvalidMimeTypeError(FileExtractionError):
    """Raised when the declared MIME type does not match the file type."""


class InvalidDocumentError(FileExtractionError):
    """Raised when a document cannot be parsed."""


class EmptyDocumentError(FileExtractionError):
    """Raised when a document contains no extractable text."""


@dataclass(frozen=True)
class ExtractedFile:
    text: str
    extension: str
    source_type: str


SOURCE_TYPES = {
    ".pdf": "pdf_upload",
    ".docx": "docx_upload",
    ".txt": "text_upload",
    ".md": "markdown_upload",
    ".tex": "latex_upload",
}

ALLOWED_MIME_TYPES = {
    ".pdf": {
        "application/pdf",
        "application/octet-stream",
    },
    ".docx": {
        (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        "application/zip",
        "application/octet-stream",
    },
    ".txt": {
        "text/plain",
        "application/octet-stream",
    },
    ".md": {
        "text/markdown",
        "text/plain",
        "application/octet-stream",
    },
    ".tex": {
        "application/x-tex",
        "text/x-tex",
        "text/plain",
        "application/octet-stream",
    },
}


def extract_text_from_file(
    filename: str,
    content: bytes,
    content_type: str | None,
) -> ExtractedFile:
    extension = Path(filename).suffix.lower()

    if extension not in SOURCE_TYPES:
        supported_types = ", ".join(SOURCE_TYPES)
        raise UnsupportedFileTypeError(
            f"Unsupported file type. Supported types: {supported_types}"
        )

    normalized_mime_type = (content_type or "").split(";")[0].strip().lower()

    if (
        normalized_mime_type
        and normalized_mime_type not in ALLOWED_MIME_TYPES[extension]
    ):
        raise InvalidMimeTypeError(
            f"Unexpected MIME type for {extension}: {normalized_mime_type}"
        )

    if extension == ".pdf":
        extracted_text = _extract_pdf_text(content)
    elif extension == ".docx":
        extracted_text = _extract_docx_text(content)
    else:
        extracted_text = _extract_plain_text(content)

    extracted_text = extracted_text.strip()

    if not extracted_text:
        raise EmptyDocumentError(
            "No readable text could be extracted. Scanned PDFs require OCR."
        )

    return ExtractedFile(
        text=extracted_text,
        extension=extension,
        source_type=SOURCE_TYPES[extension],
    )


def _extract_pdf_text(content: bytes) -> str:
    if not content.startswith(b"%PDF"):
        raise InvalidDocumentError(
            "The uploaded file does not appear to be a valid PDF"
        )

    try:
        reader = PdfReader(BytesIO(content))
        page_texts = [
            page_text
            for page in reader.pages
            if (page_text := (page.extract_text() or "").strip())
        ]
        return "\n\n".join(page_texts)
    except Exception as error:
        raise InvalidDocumentError("The PDF could not be read") from error


def _extract_docx_text(content: bytes) -> str:
    _validate_docx_container(content)

    try:
        document = Document(BytesIO(content))
        text_blocks = [
            paragraph_text
            for paragraph in document.paragraphs
            if (paragraph_text := paragraph.text.strip())
        ]

        for table in document.tables:
            for row in table.rows:
                cell_values = [
                    cell_text
                    for cell in row.cells
                    if (cell_text := cell.text.strip())
                ]
                if cell_values:
                    text_blocks.append(" | ".join(cell_values))

        return "\n\n".join(text_blocks)
    except Exception as error:
        raise InvalidDocumentError("The DOCX file could not be read") from error


def _validate_docx_container(content: bytes) -> None:
    try:
        with ZipFile(BytesIO(content)) as archive:
            file_names = set(archive.namelist())
    except BadZipFile as error:
        raise InvalidDocumentError(
            "The uploaded file is not a valid DOCX document"
        ) from error

    required_files = {"[Content_Types].xml", "word/document.xml"}

    if not required_files.issubset(file_names):
        raise InvalidDocumentError(
            "The uploaded file is not a valid DOCX document"
        )


def _extract_plain_text(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise InvalidDocumentError(
            "Text files must use UTF-8 encoding"
        ) from error
